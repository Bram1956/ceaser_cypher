ENCODE_EVEN=2
EVENER=0
STEP_INCREMENTAL=1
from docx import Document
import caser_cipher

def word_encrpyt(user_document,user_shift):
    #creation
    document_source=Document(user_document)
    document_result=Document()
    #heading


    for paragraph in  range(EVENER,len(document_source.paragraphs),STEP_INCREMENTAL):
                        if paragraph %ENCODE_EVEN==EVENER:
                            document_result.add_paragraph(document_source.paragraphs[paragraph].text)
                            document_result.save('my_python_encrypt.docx')
                        elif paragraph %ENCODE_EVEN !=EVENER:
                            document_result.add_paragraph(caser_cipher.caeser_text(document_source.paragraphs[paragraph].text,user_shift))
                            document_result.save('my_python_encrypt.docx')

user_document=input("Document to encyrpt:")
user_shift=int(input("Ceaser shift:"))

word_encrpyt(user_document,user_shift)